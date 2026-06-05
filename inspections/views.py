from .forms import HACCPFileForm
from .models import HACCPFile
# تفاصيل منشأة
from django.contrib import messages
from django.shortcuts import get_object_or_404, redirect, render

def establishment_detail(request, pk):
    establishment = get_object_or_404(Establishment, pk=pk)
    qualification = getattr(establishment.qualification_followups.first(), 'current_status', None)
    haccp_files = establishment.haccp_files.all()
    water_classifications = establishment.water_classifications.select_related('inspector').order_by('-classified_at', '-created_at')
    if request.method == 'POST':
        form = HACCPFileForm(request.POST, request.FILES)
        if form.is_valid():
            haccp_file = form.save(commit=False)
            haccp_file.establishment = establishment
            haccp_file.save()
            messages.success(request, 'تم رفع الملف بنجاح')
            return redirect('establishment_detail', pk=pk)
    else:
        form = HACCPFileForm()
    return render(request, 'inspections/establishment_detail.html', {
        'establishment': establishment,
        'qualification': qualification,
        'haccp_files': haccp_files,
        'water_classifications': water_classifications,
        'form': form,
    })
import os
import hashlib
import json
import tempfile
import zipfile
from collections import OrderedDict
from decimal import Decimal, InvalidOperation
from io import BytesIO, StringIO

from django.conf import settings
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.core.management import call_command
from django.contrib.staticfiles import finders
from django.core.cache import cache
from django.core.paginator import Paginator
from django.db.models import Avg, Count, IntegerField, OuterRef, Prefetch, Q, Subquery, Value
from django.db.models.functions import Coalesce
from django.forms import modelformset_factory
from django.http import FileResponse, HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.template.loader import get_template
from django.utils import timezone
from django.utils.http import url_has_allowed_host_and_scheme
from django.views.decorators.http import require_POST
from openpyxl import Workbook
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from xhtml2pdf import files as pisa_files
from xhtml2pdf import pisa

from .evaluation_template_data import READY_CORRECTIVE_ACTIONS_BY_CODE
from .forms import (
    CorrectiveActionForm,
    EstablishmentForm,
    EvaluationForm,
    EvaluationHeaderForm,
    EvaluationItemForm,
    EvaluationRecordCheckForm,
    EvaluationTeamMemberForm,
    QualificationFollowUpForm,
)
from .models import (
    CorrectiveActionLog,
    Criterion,
    Establishment,
    Evaluation,
    EvaluationActivityLog,
    EvaluationImage,
    EvaluationItem,
    EvaluationRecordCheck,
    EvaluationSection,
    EvaluationTeamMember,
    Governorate,
    QualificationFollowUp,
    RequiredRecord,
    UserProfile,
    WaterFactoryClassification,
    Wilayat,
)

PAGE_SIZE = 25

REFERENCE_DATA_CACHE_KEY = 'inspection_reference_data_v1'
REFERENCE_DATA_CACHE_TIMEOUT = 60 * 30
DASHBOARD_CACHE_TIMEOUT = 60

_ARABIC_TO_LATIN_DIGITS = str.maketrans('٠١٢٣٤٥٦٧٨٩', '0123456789')


def _clear_evaluation_pdf_cache(evaluation_id):
    cache.delete_many([
        f'pdf_bytes:eval:{evaluation_id}',
        f'pdf_bytes:v2:eval:{evaluation_id}',
        f'pdf_bytes:v3:eval:{evaluation_id}',
    ])


def _normalize_digit_text(value):
    if not value:
        return ''
    return str(value).translate(str.maketrans('٠١٢٣٤٥٦٧٨٩', '0123456789'))


def _normalize_criterion_code(code):
    return str(code or '').translate(_ARABIC_TO_LATIN_DIGITS).strip()


def _get_ready_corrective_action_text(criterion_code):
    normalized_code = _normalize_criterion_code(criterion_code)
    return READY_CORRECTIVE_ACTIONS_BY_CODE.get(normalized_code, '')


def _get_existing_sqlite_database_path():
    database = settings.DATABASES.get('default', {})
    engine = database.get('ENGINE', '')
    name = database.get('NAME')
    if 'sqlite3' not in engine or not name or name == ':memory:':
        return None

    db_path = os.path.abspath(os.fspath(name))
    if os.path.exists(db_path):
        return db_path
    return None


def _build_database_json_dump():
    output = StringIO()
    call_command(
        'dumpdata',
        '--natural-foreign',
        '--natural-primary',
        exclude=['contenttypes', 'auth.Permission'],
        stdout=output,
    )
    return output.getvalue().encode('utf-8')


def _get_profile(user):
    """إرجاع UserProfile للمستخدم أو None إن لم يوجد."""
    try:
        return user.userprofile
    except UserProfile.DoesNotExist:
        return None


def _get_reference_data():
    reference_data = cache.get(REFERENCE_DATA_CACHE_KEY)
    if reference_data is None:
        reference_data = {
            'governorates': list(
                Governorate.objects.order_by('name_ar').values('id', 'name_ar')
            ),
            'wilayats': list(
                Wilayat.objects.order_by('governorate__name_ar', 'name_ar').values(
                    'id', 'name_ar', 'governorate_id'
                )
            ),
        }
        cache.set(REFERENCE_DATA_CACHE_KEY, reference_data, REFERENCE_DATA_CACHE_TIMEOUT)
    return reference_data


def _get_activity_options(qs):
    return list(
        qs.exclude(activity_type='')
        .order_by('activity_type')
        .values_list('activity_type', flat=True)
        .distinct()
    )


def _with_qualification_visit_no(queryset):
    visit_no_subquery = (
        QualificationFollowUp.objects.filter(
            establishment_id=OuterRef('establishment_id'),
            pk__lte=OuterRef('pk'),
        )
        .order_by()
        .values('establishment_id')
        .annotate(total=Count('id'))
        .values('total')[:1]
    )
    return queryset.annotate(
        cached_visit_no=Coalesce(
            Subquery(visit_no_subquery, output_field=IntegerField()),
            Value(1),
        )
    )


def _build_dashboard_cache_key(user_id, governorate_id, wilayat_id, classification, approval_status, date_from, date_to):
    raw = '|'.join([
        str(user_id),
        governorate_id,
        wilayat_id,
        classification,
        approval_status,
        date_from,
        date_to,
    ])
    digest = hashlib.md5(raw.encode('utf-8')).hexdigest()
    return f'dashboard_ctx:{digest}'


def _apply_rbac(qs_establishments, qs_evaluations, profile):
    """
    تطبيق قيود الوصول المبنية على الدور:
    - admin / central / reviewer: يرون كل شيء
    - manager: يرون محافظتهم فقط
    - inspector: يرون كل المنشآت، لكن تقييماتهم فقط
    """
    if profile is None:
        return qs_establishments, qs_evaluations

    role = profile.role
    if role == 'manager' and profile.governorate_id:
        if qs_establishments is not None:
            qs_establishments = qs_establishments.filter(governorate=profile.governorate)
        if qs_evaluations is not None:
            qs_evaluations = qs_evaluations.filter(establishment__governorate=profile.governorate)
    elif role == 'inspector':
        if qs_evaluations is not None:
            qs_evaluations = qs_evaluations.filter(inspector=profile.user)

    return qs_establishments, qs_evaluations


def _get_allowed_evaluations(user, queryset=None):
    if queryset is None:
        queryset = Evaluation.objects.all()
    _, queryset = _apply_rbac(None, queryset, _get_profile(user))
    return queryset


def _build_default_corrective_action(item):
    criterion_code = item.criterion.code
    ready_text = _get_ready_corrective_action_text(criterion_code)
    if ready_text:
        return ready_text

    criterion_text = (item.criterion.text_ar or '').strip()
    risk_text = item.criterion.get_risk_level_display()
    risk_level = item.criterion.risk_level

    if risk_level == 'critical':
        timeline = 'خلال 24 ساعة'
        step_two = 'تنفيذ إجراء احتواء فوري وإيقاف الممارسة الخطرة حتى تصحيحها بالكامل.'
        step_three = 'التحقق الموقعي من الفاعلية خلال 24-48 ساعة مع اعتماد المراجع الفني.'
    elif risk_level == 'high':
        timeline = 'خلال 3 أيام'
        step_two = 'تنفيذ إجراء تصحيحي عاجل يمنع تكرار عدم الاستيفاء.'
        step_three = 'التحقق من الفاعلية خلال 3-5 أيام وتوثيق نتائج المتابعة.'
    elif risk_level == 'medium':
        timeline = 'خلال 7 أيام'
        step_two = 'تنفيذ الإجراء التصحيحي وفق خطة عمل محددة بالمسؤوليات.'
        step_three = 'التحقق من الفاعلية خلال 7-10 أيام وتحديث السجلات.'
    else:
        timeline = 'خلال 14 يومًا'
        step_two = 'تنفيذ تحسين تشغيلي ومعالجة سبب عدم الاستيفاء.'
        step_three = 'التحقق الدوري من الفاعلية وتوثيق الإغلاق النهائي.'

    return (
        f'1) معالجة عدم الاستيفاء في البند {criterion_code} ({criterion_text}) وإزالة السبب الجذري.\n'
        f'2) {step_two}\n'
        f'3) {step_three}\n'
        f'4) تاريخ الإغلاق المستهدف: {timeline}.\n'
        f'مستوى الخطورة: {risk_text}.'
    )


def _sync_corrective_actions_for_evaluation(evaluation, user, pre_loaded_items=None):
    if pre_loaded_items is not None:
        items = pre_loaded_items
    else:
        items = list(
            EvaluationItem.objects.filter(evaluation=evaluation)
            .select_related('criterion')
        )
    existing_logs = {
        log.criterion_id: log
        for log in CorrectiveActionLog.objects.filter(evaluation=evaluation, criterion__isnull=False)
    }

    active_criterion_ids = set()
    for item in items:
        has_corrective_action = bool((item.corrective_action or '').strip())
        if item.status != 'non_compliant' or not has_corrective_action:
            continue

        active_criterion_ids.add(item.criterion_id)
        title = f'إجراء تصحيحي للبند {item.criterion.code}'
        details = item.corrective_action.strip()
        log = existing_logs.get(item.criterion_id)

        if log is None:
            CorrectiveActionLog.objects.create(
                evaluation=evaluation,
                criterion=item.criterion,
                created_by=user,
                title=title,
                details=details,
                assigned_to='يحدد لاحقاً',
                status='open',
            )
            continue

        updated_fields = []
        if log.title != title:
            log.title = title
            updated_fields.append('title')
        if log.details != details:
            log.details = details
            updated_fields.append('details')
        if not log.assigned_to:
            log.assigned_to = 'يحدد لاحقاً'
            updated_fields.append('assigned_to')
        if updated_fields:
            log.save(update_fields=updated_fields)

    stale_logs = [
        log.pk for criterion_id, log in existing_logs.items()
        if criterion_id not in active_criterion_ids and log.status == 'open'
    ]
    if stale_logs:
        CorrectiveActionLog.objects.filter(pk__in=stale_logs).delete()


ISIC4_FOOD_ACTIVITIES = [
    ('1010', 'تجهيز وحفظ اللحوم'),
    ('1020', 'تجهيز وحفظ الأسماك والقشريات والرخويات'),
    ('1030', 'تجهيز وحفظ الفواكه والخضروات'),
    ('1040', 'صناعة الزيوت والدهون النباتية والحيوانية'),
    ('1050', 'صناعة منتجات الألبان'),
    ('1061', 'صناعة منتجات مطاحن الحبوب'),
    ('1062', 'صناعة النشا ومنتجاته'),
    ('1071', 'صناعة منتجات المخابز'),
    ('1072', 'صناعة السكر'),
    ('1073', 'صناعة الكاكاو والشوكولاتة والحلويات السكرية'),
    ('1074', 'صناعة المعكرونة والشعيرية والكسكس ومنتجات النشا المماثلة'),
    ('1075', 'صناعة الوجبات والأطباق الجاهزة'),
    ('1079', 'صناعة منتجات غذائية أخرى غير مصنفة في موضع آخر'),
    ('1080', 'صناعة الأعلاف الحيوانية المحضرة'),
    ('1101', 'تقطير ومزج المشروبات الروحية'),
    ('1102', 'صناعة النبيذ'),
    ('1103', 'صناعة مشروبات الشعير والملت'),
    ('1104', 'صناعة المشروبات الغازية وإنتاج المياه المعدنية والمياه المعبأة الأخرى'),
    ('014102', 'انتاج اللبن البقر ي الخام من البقر أو الجاموس'),
    ('014303', 'انتاج لبن الأبل الخام'),
    ('014402', 'انتاج لبن الضأن أو الماعز الخام'),
    ('014603', 'انتاج البيض'),
    ('101004', 'صنع منتجات اللحوم (سجق، برجر، مرتديلا، دقيق، مسحوق، مجفف ... إلخ)'),
    ('101005', 'تذويب الشحوم'),
    ('101006', 'استخلاص وتكرير الدهون الحيوانية'),
    ('101011', 'تجهيز وتغليف البيض'),
    ('103001', 'تحضير أو تعليب أو حفظ الفواكه وعصيرها'),
    ('103002', 'تحضير أو تعليب أو حفظ الخضراوات الطازجة أو المطبوخة وعصيرها'),
    ('103004', 'تجفيف وتعبئة التمور والعنب والتين وصنع منتجاتها'),
    ('103005', 'صنع صلصة الطماطم'),
    ('103006', 'صنع المربى والمرملاد وهلامات المائدة (الجيلي)'),
    ('103007', 'صنع المخللات والطرشي (التخليل)'),
    ('103008', 'حفظ وتحميص وتعبئة المكسرات بأنواعها'),
    ('103010', 'صنع رقائق البطاطس'),
    ('103011', 'تجفيف الخضر البقولية بالوسائل الصناعية'),
    ('103012', 'صنع التوفو (عجينة البقول)'),
    ('104001', 'انتاج وتكرير زيت السمسم'),
    ('104002', 'انتاج وتكرير الزيوت النباتية الأخرى'),
    ('104003', 'استخلاص وتكرير الزيوت من الأسماك والأحياء البحرية'),
    ('104004', 'صنع المارجرين ومخلفات عصر الزيوت والدهون (الحيوانية والنباتية)'),
    ('104005', 'تعبئة الزيوت والدهون النباتية بعد تكريرها'),
    ('104006', 'صنع الدهون الحيوانية الصالحة للأكل'),
    ('104007', 'انتاج وتكرير زيت الزيتون'),
    ('104008', 'انتاج وتكرير زيت عباد الشمس'),
    ('104009', 'صنع واستخلاص مادة الكركمين والزيوت الاساسية من الكركم'),
    ('105001', 'تجهيز وبسترة وتعبئة الحليب واللبن الطازج'),
    ('105002', 'انتاج الروب (الزبادي)، القشدة، الزبدة، السمن ... إلخ'),
    ('105003', 'انتاج أنواع الأجبان (أبيض، جاف، مطبوخ ... إلخ)'),
    ('105004', 'صنع الحليب المجفف (البودرة) والمكثف'),
    ('105005', 'صنع المثلجات (الآيس كريم)'),
    ('105006', 'صنع اللاكتوز والكاريين'),
    ('106101', 'طحن وتعبئة الحبوب (قمح، ذرة، شعير ... إلخ)'),
    ('106102', 'طحن وتعبئة الأرز وتبييضه'),
    ('106103', 'انتاج دقيق من الأرز'),
    ('106104', 'طحن الخضراوات البقولية المجففة والجوزيات المعدة للأكل وتعبئتها'),
    ('106105', 'صنع الدقيق والعجين للمخابز'),
    ('106106', 'صنع أطعمة الإفطار (كورن فليكس، شيبس ... إلخ)، من الحبوب على شكل رقائق'),
    ('106201', 'صنع النشاء من الذرة أو الأرز أو الحبوب الأخرى أو البطاطا أو من النباتات الأخرى'),
    ('106202', 'طحن الذرة المنداه'),
    ('106203', 'انتاج وتكرير زيت الذرة'),
    ('106204', 'صنع الجلوكوز وشراب الجلوكوز وسكر الشعير (الملتوز)'),
    ('107104', 'صنع البسكويت بأنواعه'),
    ('107201', 'صناعة السكر الخام من قصب السكر والبنجر'),
    ('107202', 'تكرير السكر الخام'),
    ('107203', 'تعبئة السكر'),
    ('107301', 'صنع الشوكولاتة والحلويات المصنوعة من الشوكولاتة'),
    ('107302', 'صنع الحلويات السكرية'),
    ('107303', 'صنع العلك (اللبان)'),
    ('107304', 'صنع الحلاوة الطحينية'),
    ('107306', 'صنع الكاكاو على شكل عجينه أو ألواح أو مسحوق أو زبدة أو دهن أو زيت الكاكاو'),
    ('107401', 'صنع المعكرونة بأنواعها'),
    ('107402', 'صنع العجائن المحشوة المعبأة في علب أو المجمدة'),
    ('107403', 'صنع العجائن المحشوة المطبوخة أو غير المطبوخة'),
    ('107500', 'صنع وجبات وأطباق جاهزة'),
    ('107903', 'تكرير وطحن ملح الطعام'),
    ('107904', 'صنع الحساء بكافة أشكاله، وصنع التوابل ومرق التوابل والخل والخميرة'),
    ('107905', 'صنع العسل الأسود (الدبس) وعسل السكر'),
    ('107906', 'صنع خلاصات ومكسبات الطعم للمواد الغذائية والمشروبات (ماء الورد، الفانيليا ... إلخ)'),
    ('110302', 'صنع مشروبات الشعير غير الكحولية'),
    ('110401', 'صنع المشروبات الغازية المرطبة'),
    ('110402', 'صنع المشروبات المنكهة بخلاصات أو أرواح الفاكهة'),
    ('110403', 'انتاج وتعبئة مياه الشرب المعبأة في عبوات'),
    ('110405', 'انتاج وتعبئة المياه المعدنية'),
    ('107101', 'صنع الخبز ومنتجاته'),
    ('107102', 'صنع الفطائر بأنواعها'),
    ('107103', 'صنع الجاتوه والبيتيفور والكعك بأنواعه'),
    ('107105', 'صنع خبز الأكلات الشعبية ورقائق المعجنات النيئة'),
    ('107107', 'صنع الحلويات العربية تشمل المعمول وخلافه'),
    ('107109', 'صنع المتاي'),
    ('107901', 'تحميص البن أو طحنه أو تعبئته وصنع بدائل البن'),
    ('107907', 'خلط الشاي وتعبئته'),
    ('107908', 'طحن وتعبئة البهارات والتوابل'),
    ('110404', 'صنع المنتجات الحرفية بتقطير الزهور والاعشاب'),
    ('110406', 'صنع المنتجات الحرفية لإنتاج ماء وزيت اللبان'),
    ('101003', 'تبريد وتجميد اللحوم'),
    ('103003', 'تبريد وتجميد الفواكه والخضراوات'),
]


def home(request):
    return render(request, 'inspections/home.html', {'minimal_nav': True})


def _get_allowed_establishments(user, queryset=None):
    if queryset is None:
        queryset = Establishment.objects.all()
    queryset, _ = _apply_rbac(queryset, None, _get_profile(user))
    return queryset


def _get_allowed_water_classifications(user, queryset=None):
    if queryset is None:
        queryset = WaterFactoryClassification.objects.all()
    profile = _get_profile(user)
    if profile and profile.role == 'manager' and profile.governorate_id:
        queryset = queryset.filter(establishment__governorate=profile.governorate)
    elif profile and profile.role == 'inspector':
        queryset = queryset.filter(inspector=profile.user)
    return queryset


@login_required
def water_factory_classification(request):
    classification_levels = [
        {'grade': 'A+', 'range': '95-100%', 'risk': 'منخفض جداً', 'decision': 'اعتماد كامل'},
        {'grade': 'A', 'range': '90-94%', 'risk': 'منخفض', 'decision': 'مؤهل'},
        {'grade': 'B', 'range': '80-89%', 'risk': 'متوسط', 'decision': 'خطة تحسين'},
        {'grade': 'C', 'range': '70-79%', 'risk': 'مرتفع', 'decision': 'إعادة تقييم'},
        {'grade': 'D', 'range': 'أقل من 70%', 'risk': 'خطر عالي', 'decision': 'إيقاف أو تعليق'},
    ]
    risk_weights = [
        {'level': 'حرج جداً', 'score': '10 نقاط'},
        {'level': 'عالي', 'score': '7 نقاط'},
        {'level': 'متوسط', 'score': '4 نقاط'},
        {'level': 'منخفض', 'score': '2 نقطة'},
    ]
    assessment_sections = [
        'البيانات الأساسية للمصنع',
        'الموقع العام والمبنى الخارجي',
        'مصدر المياه الخام',
        'نظام معالجة المياه',
        'التعبئة والإنتاج',
        'النظافة والتطهير',
        'مكافحة الآفات',
        'النظافة الشخصية والعاملين',
        'المختبر وضبط الجودة',
        'المواصفات الكيميائية والميكروبية للمنتج النهائي',
        'نظام HACCP',
        'التخزين والنقل',
        'التتبع وسحب المنتج',
        'الوثائق والسجلات',
    ]
    critical_items = [
        'وجود E. coli بالمياه',
        'فشل التعقيم النهائي',
        'تلوث المنتج النهائي',
        'عدم صلاحية مصدر المياه',
        'فشل نظام RO',
        'وجود تسربات أو صدأ بالخزانات',
        'عدم وجود برنامج تنظيف',
        'استخدام مواد كيميائية غير معتمدة',
        'غياب نظام HACCP',
        'عدم وجود تتبع للمنتج',
    ]
    context = {
        'classification_levels': classification_levels,
        'risk_weights': risk_weights,
        'assessment_sections': assessment_sections,
        'critical_items': critical_items,
    }
    water_qs = _get_allowed_water_classifications(
        request.user,
        WaterFactoryClassification.objects.select_related(
            'establishment',
            'establishment__governorate',
            'establishment__wilayat',
            'inspector',
        ),
    )
    allowed_establishments = _get_allowed_establishments(request.user)
    grade_counts = {
        row['grade']: row['total']
        for row in water_qs.values('grade').annotate(total=Count('id'))
    }
    total_water_classifications = water_qs.count()
    grade_chart_items = [
        {
            'grade': item['grade'],
            'label': item['range'],
            'count': grade_counts.get(item['grade'], 0),
            'pct': round((grade_counts.get(item['grade'], 0) / total_water_classifications) * 100, 1) if total_water_classifications else 0,
        }
        for item in classification_levels
    ]
    context.update({
        'total_water_classifications': total_water_classifications,
        'classified_factories_count': water_qs.values('establishment_id').distinct().count(),
        'available_factories_count': allowed_establishments.filter(status='active').count(),
        'avg_water_percentage': water_qs.aggregate(avg=Avg('percentage'))['avg'] or 0,
        'latest_water_classifications': water_qs.order_by('-classified_at', '-created_at')[:10],
        'grade_counts': grade_counts,
        'grade_chart_items': grade_chart_items,
        'hide_establishments_nav': True,
        'hide_evaluations_nav': True,
        'hide_nav_actions': True,
    })
    return render(request, 'inspections/water_factory_classification.html', context)


@login_required
def water_factory_evaluation_form(request):
    factories = (
        _get_allowed_establishments(
            request.user,
            Establishment.objects.select_related('governorate', 'wilayat'),
        )
        .filter(status='active')
        .only(
            'id',
            'establishment_no',
            'commercial_name',
            'activity_type',
            'license_no',
            'commercial_reg',
            'manager_name',
            'contact_phone',
            'governorate__name_ar',
            'wilayat__name_ar',
        )
        .order_by('commercial_name', 'id')
    )

    if request.method == 'POST':
        factory_id = request.POST.get('factory_id')
        if not factory_id:
            messages.error(request, 'يرجى اختيار المصنع قبل حفظ التصنيف.')
            return redirect('water_factory_evaluation_form')

        establishment = get_object_or_404(factories, pk=factory_id)
        try:
            total_possible = Decimal(request.POST.get('total_possible_points', '0') or '0')
            total_earned = Decimal(request.POST.get('total_earned_points', '0') or '0')
            percentage = Decimal(request.POST.get('percentage', '0') or '0')
        except InvalidOperation:
            messages.error(request, 'تعذر حفظ التصنيف بسبب قيمة رقمية غير صحيحة.')
            return redirect('water_factory_evaluation_form')

        grade = request.POST.get('grade', '').strip()
        valid_grades = {value for value, _ in WaterFactoryClassification.GRADE_CHOICES}
        if grade not in valid_grades:
            messages.error(request, 'تعذر حفظ التصنيف بسبب تصنيف نهائي غير صحيح.')
            return redirect('water_factory_evaluation_form')

        try:
            critical_count = int(request.POST.get('critical_count', '0') or 0)
        except ValueError:
            critical_count = 0

        try:
            items_payload = json.loads(request.POST.get('items_payload', '[]') or '[]')
            if not isinstance(items_payload, list):
                items_payload = []
        except json.JSONDecodeError:
            items_payload = []

        WaterFactoryClassification.objects.create(
            establishment=establishment,
            inspector=request.user,
            total_possible_points=total_possible,
            total_earned_points=total_earned,
            percentage=percentage,
            grade=grade,
            decision=request.POST.get('decision', '').strip() or '-',
            critical_count=max(0, critical_count),
            items_payload=items_payload,
        )
        messages.success(
            request,
            f'تم حفظ تصنيف مصنع المياه {establishment.commercial_name}. التصنيف: {grade} - النسبة: {percentage}%',
        )
        return redirect('water_factory_classification')

    return render(
        request,
        'inspections/water_factory_evaluation_form.html',
        {
            'factories': factories,
        },
    )


def user_login(request):
    next_url = (request.GET.get('next') or request.POST.get('next') or '').strip()

    if request.user.is_authenticated:
        if next_url and url_has_allowed_host_and_scheme(
            url=next_url,
            allowed_hosts={request.get_host()},
            require_https=request.is_secure(),
        ):
            return redirect(next_url)
        return redirect('dashboard')

    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        password = request.POST.get('password', '')

        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            if next_url and url_has_allowed_host_and_scheme(
                url=next_url,
                allowed_hosts={request.get_host()},
                require_https=request.is_secure(),
            ):
                return redirect(next_url)
            return redirect('dashboard')

        messages.error(request, 'اسم المستخدم أو كلمة السر غير صحيحة.')

    context = {
        'minimal_nav': True,
        'next': next_url,
    }
    return render(request, 'inspections/login.html', context)


@login_required
def user_logout(request):
    logout(request)
    return redirect('home')


@login_required
def qualification_followup_list(request):
    qs = _with_qualification_visit_no(
        QualificationFollowUp.objects.select_related('establishment', 'evaluation').all()
    )

    q = request.GET.get('q', '').strip()
    governorate = request.GET.get('governorate', '').strip()
    status = request.GET.get('status', '').strip()
    activity = request.GET.get('activity', '').strip()

    if q:
        normalized_q = _normalize_digit_text(q)
        filters = (
            Q(establishment_name__icontains=q) |
            Q(activity_type__icontains=q) |
            Q(governorate__icontains=q) |
            Q(establishment__license_no__icontains=q)
        )
        if normalized_q.isdigit():
            filters |= Q(establishment__establishment_no=int(normalized_q))
        qs = qs.filter(filters)
    if governorate:
        qs = qs.filter(governorate=governorate)
    if status:
        qs = qs.filter(current_status=status)
    if activity:
        qs = qs.filter(activity_type=activity)

    if request.method == 'POST':
        form = QualificationFollowUpForm(request.POST)
        if form.is_valid():
            obj = form.save()
            messages.success(request, f'تم حفظ متابعة التأهيل للمنشأة: {obj.establishment_name}')
            return redirect('qualification_followup_list')
    else:
        form = QualificationFollowUpForm()

    today = timezone.localdate()
    stats = qs.aggregate(
        total=Count('id'),
        avg_progress=Avg('progress_percent'),
        completed=Count('id', filter=Q(current_status='completed')),
        in_progress=Count('id', filter=Q(current_status='in_progress')),
        stalled=Count('id', filter=Q(current_status='stalled')),
        overdue=Count(
            'id',
            filter=Q(expected_completion_date__lt=today) & ~Q(current_status='completed'),
        ),
    )
    total_count = stats['total'] or 0
    completed_count = stats['completed'] or 0
    in_progress_count = stats['in_progress'] or 0
    stalled_count = stats['stalled'] or 0
    overdue_count = stats['overdue'] or 0
    avg_progress = round(stats['avg_progress'] or 0, 1)

    paginator = Paginator(qs, PAGE_SIZE)
    page_obj = paginator.get_page(request.GET.get('page'))
    reference_data = _get_reference_data()
    activity_options = (
        QualificationFollowUp.objects
        .exclude(activity_type='')
        .order_by('activity_type')
        .values_list('activity_type', flat=True)
        .distinct()
    )
    establishment_options = Establishment.objects.select_related('governorate').only(
        'id', 'establishment_no', 'commercial_name', 'activity_type', 'governorate__name_ar'
    ).order_by('establishment_no', 'commercial_name')

    return render(request, 'inspections/qualification_followup_list.html', {
        'form': form,
        'items': page_obj,
        'page_obj': page_obj,
        'q': q,
        'governorates': reference_data['governorates'],
        'selected_governorate': governorate,
        'selected_status': status,
        'selected_activity': activity,
        'status_choices': QualificationFollowUp.STATUS_CHOICES,
        'activity_options': activity_options,
        'total_count': total_count,
        'completed_count': completed_count,
        'in_progress_count': in_progress_count,
        'stalled_count': stalled_count,
        'overdue_count': overdue_count,
        'avg_progress': avg_progress,
        'establishment_options': establishment_options,
    })


@login_required
def dashboard(request):
    profile = _get_profile(request.user)

    governorate_id = request.GET.get('governorate', '').strip()
    wilayat_id = request.GET.get('wilayat', '').strip()
    classification = request.GET.get('classification', '').strip()
    approval_status = request.GET.get('approval_status', '').strip()
    date_from = request.GET.get('date_from', '').strip()
    date_to = request.GET.get('date_to', '').strip()

    cache_key = _build_dashboard_cache_key(
        request.user.id,
        governorate_id,
        wilayat_id,
        classification,
        approval_status,
        date_from,
        date_to,
    )
    cached_context = cache.get(cache_key)
    if cached_context is not None:
        return render(request, 'inspections/dashboard.html', cached_context)

    establishments_qs = Establishment.objects.all()
    evaluations_qs = Evaluation.objects.select_related('establishment')
    establishments_qs, evaluations_qs = _apply_rbac(establishments_qs, evaluations_qs, profile)

    if governorate_id:
        establishments_qs = establishments_qs.filter(governorate_id=governorate_id)
        evaluations_qs = evaluations_qs.filter(establishment__governorate_id=governorate_id)
    if wilayat_id:
        establishments_qs = establishments_qs.filter(wilayat_id=wilayat_id)
        evaluations_qs = evaluations_qs.filter(establishment__wilayat_id=wilayat_id)
    if classification:
        evaluations_qs = evaluations_qs.filter(classification=classification)
    if approval_status:
        evaluations_qs = evaluations_qs.filter(approval_status=approval_status)
    if date_from:
        evaluations_qs = evaluations_qs.filter(visit_date__gte=date_from)
    if date_to:
        evaluations_qs = evaluations_qs.filter(visit_date__lte=date_to)

    total_establishments = establishments_qs.count()
    evaluation_stats = evaluations_qs.aggregate(
        total=Count('id'),
        completed=Count('id', filter=Q(approval_status='completed')),
        avg_percentage=Avg('percentage'),
    )
    total_evaluations = evaluation_stats['total'] or 0
    completed_evaluations = evaluation_stats['completed'] or 0
    open_actions = CorrectiveActionLog.objects.exclude(status='closed').filter(evaluation__in=evaluations_qs).count()
    non_compliant_items = EvaluationItem.objects.filter(
        status='non_compliant',
        criterion__is_active=True,
        evaluation__in=evaluations_qs,
    ).count()
    avg_percentage = evaluation_stats['avg_percentage'] or 0

    by_governorate = list(
        establishments_qs.values('governorate__name_ar')
        .annotate(total=Count('id'))
        .order_by('governorate__name_ar')
    )
    recent_evaluations = list(
        evaluations_qs.only(
            'id',
            'visit_date',
            'percentage',
            'approval_status',
            'establishment__commercial_name',
        )
        .order_by('-visit_date', '-created_at')[:10]
    )
    classification_summary = list(
        evaluations_qs.values('classification').annotate(total=Count('id')).order_by('-total')
    )
    weakest_evaluations = list(
        evaluations_qs.only(
            'id',
            'visit_date',
            'percentage',
            'classification',
            'establishment__commercial_name',
        )
        .order_by('percentage', '-visit_date', '-created_at')[:5]
    )
    establishment_growth = list(
        establishments_qs.values('governorate__name_ar')
        .annotate(
            active_total=Count('id', filter=Q(status='active')),
            total=Count('id'),
        )
        .order_by('-total')[:6]
    )

    classification_labels = dict(Evaluation.CLASSIFICATION_CHOICES)
    for row in classification_summary:
        row['label'] = classification_labels.get(row['classification'], row['classification'])
        row['pct'] = round((row['total'] / total_evaluations) * 100, 1) if total_evaluations else 0

    for row in establishment_growth:
        row['active_pct'] = round((row['active_total'] / row['total']) * 100, 1) if row['total'] else 0

    STATUS_META = {
        'excellent': {'label': 'ممتاز', 'range': '86%-100%', 'description': 'مستوفي للحصول على شهادة ضبط الجودة', 'color': 'success', 'icon': 'fa-star', 'order': 1},
        'good': {'label': 'جيد', 'range': '70%-85%', 'description': 'مستوفي مع وجود فرص للتحسين', 'color': 'info', 'icon': 'fa-thumbs-up', 'order': 2},
        'acceptable': {'label': 'مقبول', 'range': '41%-69%', 'description': 'يحتاج تأهيل ومزيد من التحسين', 'color': 'warning', 'icon': 'fa-triangle-exclamation', 'order': 3},
        'weak': {'label': 'ضعيف', 'range': '0%-40%', 'description': 'إيقاف الإنتاج', 'color': 'danger', 'icon': 'fa-circle-xmark', 'order': 4},
    }

    latest_eval_subq = evaluations_qs.filter(
        establishment=OuterRef('pk')
    ).order_by('-visit_date', '-created_at')
    establishments_with_latest = establishments_qs.annotate(
        latest_cls=Subquery(latest_eval_subq.values('classification')[:1])
    )
    counted = {
        row['latest_cls']: row['total']
        for row in establishments_with_latest.values('latest_cls').annotate(total=Count('id')).order_by()
        if row['latest_cls']
    }
    no_eval_count = establishments_with_latest.filter(latest_cls__isnull=True).count()
    establishment_status_summary = []
    for key, meta in sorted(STATUS_META.items(), key=lambda x: x[1]['order']):
        count = counted.get(key, 0)
        establishment_status_summary.append({
            **meta,
            'key': key,
            'count': count,
            'pct': round((count / total_establishments) * 100, 1) if total_establishments else 0,
        })

    reference_data = _get_reference_data()

    context = {
        'total_establishments': total_establishments,
        'total_evaluations': total_evaluations,
        'completed_evaluations': completed_evaluations,
        'open_actions': open_actions,
        'non_compliant_items': non_compliant_items,
        'by_governorate': by_governorate,
        'recent_evaluations': recent_evaluations,
        'avg_percentage': round(avg_percentage, 2),
        'classification_summary': classification_summary,
        'weakest_evaluations': weakest_evaluations,
        'establishment_growth': establishment_growth,
        'establishment_status_summary': establishment_status_summary,
        'no_eval_count': no_eval_count,
        'governorates': reference_data['governorates'],
        'wilayats': reference_data['wilayats'],
        'selected_governorate': governorate_id,
        'selected_wilayat': wilayat_id,
        'selected_classification': classification,
        'selected_approval_status': approval_status,
        'date_from': date_from,
        'date_to': date_to,
        'classification_choices': Evaluation.CLASSIFICATION_CHOICES,
        'approval_status_choices': Evaluation.APPROVAL_STATUS_CHOICES,
    }

    cache.set(cache_key, context, DASHBOARD_CACHE_TIMEOUT)

    return render(request, 'inspections/dashboard.html', context)


def _build_establishment_list_context(request):
    profile = _get_profile(request.user)
    qs = (
        Establishment.objects.select_related('governorate', 'wilayat')
        .only(
            'id',
            'establishment_no',
            'commercial_name',
            'activity_type',
            'license_no',
            'direct_location_url',
            'status',
            'created_at',
            'governorate__name_ar',
            'wilayat__name_ar',
        )
        .order_by('commercial_name', 'id')
    )
    qs, _ = _apply_rbac(qs, None, profile)
    activity_options = _get_activity_options(qs)

    q = request.GET.get('q', '').strip()
    normalized_q = _normalize_digit_text(q)
    governorate_id = request.GET.get('governorate', '').strip()
    wilayat_id = request.GET.get('wilayat', '').strip()
    activity = request.GET.get('activity', '').strip()
    if q:
        filters = (
            Q(commercial_name__icontains=q) |
            Q(activity_type__icontains=q) |
            Q(license_no__icontains=q) |
            Q(commercial_reg__icontains=q)
        )
        if normalized_q.isdigit():
            filters |= Q(establishment_no=int(normalized_q))
        qs = qs.filter(filters)
    if governorate_id:
        qs = qs.filter(governorate_id=governorate_id)
    if wilayat_id:
        qs = qs.filter(wilayat_id=wilayat_id)
    if activity:
        qs = qs.filter(activity_type=activity)
    reference_data = _get_reference_data()
    total_establishments = qs.count()
    status_counts = {
        row['status']: row['total']
        for row in qs.values('status').annotate(total=Count('id'))
    }
    establishment_status_chart = [
        {
            'label': label,
            'count': status_counts.get(value, 0),
            'pct': round((status_counts.get(value, 0) / total_establishments) * 100, 1) if total_establishments else 0,
            'color': {
                'active': 'success',
                'suspended': 'warning',
                'closed': 'danger',
            }.get(value, 'secondary'),
        }
        for value, label in Establishment.STATUS_CHOICES
    ]
    establishment_governorate_chart = [
        {
            'label': row['governorate__name_ar'] or '-',
            'count': row['total'],
            'pct': round((row['total'] / total_establishments) * 100, 1) if total_establishments else 0,
        }
        for row in qs.values('governorate__name_ar').annotate(total=Count('id')).order_by('-total')[:6]
    ]

    paginator = Paginator(qs, PAGE_SIZE)
    page_obj = paginator.get_page(request.GET.get('page'))
    query_params = request.GET.copy()
    query_params.pop('page', None)

    return {
        'establishments': page_obj,
        'page_obj': page_obj,
        'q': q,
        'governorates': reference_data['governorates'],
        'wilayats': reference_data['wilayats'],
        'activity_options': activity_options,
        'selected_governorate': governorate_id,
        'selected_wilayat': wilayat_id,
        'selected_activity': activity,
        'list_querystring': query_params.urlencode(),
        'total_establishments': total_establishments,
        'active_establishments_count': status_counts.get('active', 0),
        'suspended_establishments_count': status_counts.get('suspended', 0),
        'closed_establishments_count': status_counts.get('closed', 0),
        'establishment_activity_count': qs.exclude(activity_type='').values('activity_type').distinct().count(),
        'establishment_status_chart': establishment_status_chart,
        'establishment_governorate_chart': establishment_governorate_chart,
    }


@login_required
def establishment_list(request):
    return render(
        request,
        'inspections/establishment_list.html',
        _build_establishment_list_context(request),
    )


@login_required
def establishment_create(request):
    next_url = (request.GET.get('next') or request.POST.get('next') or '').strip()
    form = EstablishmentForm(request.POST or None, request.FILES or None)
    if form.is_valid():
        obj = form.save()
        messages.success(request, f'تم حفظ المنشأة بنجاح. رقم المنشأة: {obj.establishment_no}')
        if next_url and url_has_allowed_host_and_scheme(
            url=next_url,
            allowed_hosts={request.get_host()},
            require_https=request.is_secure(),
        ):
            return redirect(next_url)
        return redirect('establishment_create')
    reference_data = _get_reference_data()
    context = _build_establishment_list_context(request)
    context.update({
        'form': form,
        'title': 'إضافة منشأة',
        'wilayats': reference_data['wilayats'],
        'isic_activities': ISIC4_FOOD_ACTIVITIES,
        'next': next_url,
        'hide_evaluations_nav': True,
        'hide_water_nav': True,
        'hide_nav_actions': True,
    })
    return render(request, 'inspections/establishment_form.html', context)


@login_required
def download_database_backup(request):
    if request.user.is_superuser:
        timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
        db_path = _get_existing_sqlite_database_path()
        if db_path:
            backup_name = f'food_safety_backup_{timestamp}.sqlite3'

            return FileResponse(
                open(db_path, 'rb'),
                as_attachment=True,
                filename=backup_name,
                content_type='application/x-sqlite3',
            )

        backup_name = f'food_safety_backup_{timestamp}.json'
        response = HttpResponse(
            _build_database_json_dump(),
            content_type='application/json; charset=utf-8',
        )
        response['Content-Disposition'] = f'attachment; filename="{backup_name}"'
        return response

    if not request.user.is_superuser:
        messages.error(request, 'غير مسموح لك بتنزيل النسخة الاحتياطية.')
        return redirect('home')


@login_required
def download_reports_backup(request):
    if request.user.is_superuser:
        timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
        backup_name = f'food_safety_reports_backup_{timestamp}.zip'
        media_root = getattr(settings, 'MEDIA_ROOT', '')
        db_path = _get_existing_sqlite_database_path()

        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', compression=zipfile.ZIP_DEFLATED) as backup_zip:
            if db_path:
                backup_zip.write(db_path, arcname=f'database/{os.path.basename(db_path)}')
            else:
                backup_zip.writestr('database/database_dump.json', _build_database_json_dump())

            if media_root:
                for image in EvaluationImage.objects.only('image').iterator():
                    if not image.image:
                        continue
                    file_name = image.image.name
                    file_path = os.path.join(media_root, file_name)
                    if os.path.exists(file_path):
                        backup_zip.write(file_path, arcname=f'media/{file_name}')

        zip_buffer.seek(0)
        return FileResponse(
            zip_buffer,
            as_attachment=True,
            filename=backup_name,
            content_type='application/zip',
        )

    if not request.user.is_superuser:
        messages.error(request, 'غير مسموح لك بتنزيل النسخة الاحتياطية.')
        return redirect('home')


@login_required
@require_POST
def evaluation_delete(request, pk):
    allowed_qs = _get_allowed_evaluations(request.user)
    evaluation = get_object_or_404(allowed_qs, pk=pk)
    evaluation_id = evaluation.pk

    establishment_name = evaluation.establishment.commercial_name
    visit_date = evaluation.visit_date

    for image in evaluation.images.only('image').iterator():
        if image.image:
            image.image.delete(save=False)

    evaluation.delete()
    _clear_evaluation_pdf_cache(evaluation_id)
    messages.success(
        request,
        f'تم حذف التقرير للمنشأة {establishment_name} بتاريخ {visit_date:%Y/%m/%d}.',
    )
    return redirect('evaluation_list')


@login_required
def evaluation_list(request):
    profile = _get_profile(request.user)
    qs = (
        Evaluation.objects.select_related(
            'establishment',
            'establishment__governorate',
            'establishment__wilayat',
        )
        .only(
            'id',
            'visit_date',
            'percentage',
            'classification',
            'approval_status',
            'establishment__commercial_name',
            'establishment__activity_type',
            'establishment__license_no',
            'establishment__establishment_no',
            'establishment__commercial_reg',
            'establishment__governorate__name_ar',
            'establishment__wilayat__name_ar',
        )
        .order_by('-visit_date', '-created_at')
    )
    _, qs = _apply_rbac(None, qs, profile)
    activity_options = _get_activity_options(
        Establishment.objects.filter(evaluations__in=qs).distinct()
    )

    governorate_id = request.GET.get('governorate', '').strip()
    wilayat_id = request.GET.get('wilayat', '').strip()
    classification = request.GET.get('classification', '').strip()
    activity = request.GET.get('activity', '').strip()
    q = request.GET.get('q', '').strip()
    normalized_q = _normalize_digit_text(q)

    if governorate_id:
        qs = qs.filter(establishment__governorate_id=governorate_id)
    if wilayat_id:
        qs = qs.filter(establishment__wilayat_id=wilayat_id)
    if classification:
        qs = qs.filter(classification=classification)
    if activity:
        qs = qs.filter(establishment__activity_type=activity)
    if q:
        filters = (
            Q(establishment__commercial_name__icontains=q) |
            Q(establishment__activity_type__icontains=q) |
            Q(establishment__license_no__icontains=q) |
            Q(establishment__commercial_reg__icontains=q)
        )
        if normalized_q.isdigit():
            filters |= Q(establishment__establishment_no=int(normalized_q))
        qs = qs.filter(filters)

    total_evaluations = qs.count()
    approval_counts = {
        row['approval_status']: row['total']
        for row in qs.values('approval_status').annotate(total=Count('id'))
    }
    classification_counts = {
        row['classification']: row['total']
        for row in qs.values('classification').annotate(total=Count('id'))
    }
    classification_labels = dict(Evaluation.CLASSIFICATION_CHOICES)
    evaluation_classification_chart = [
        {
            'label': classification_labels.get(value, value),
            'count': classification_counts.get(value, 0),
            'pct': round((classification_counts.get(value, 0) / total_evaluations) * 100, 1) if total_evaluations else 0,
            'color': {
                'excellent': 'success',
                'good': 'primary',
                'acceptable': 'warning',
                'weak': 'danger',
            }.get(value, 'secondary'),
        }
        for value, _ in Evaluation.CLASSIFICATION_CHOICES
    ]
    evaluation_governorate_chart = [
        {
            'label': row['establishment__governorate__name_ar'] or '-',
            'count': row['total'],
            'pct': round((row['total'] / total_evaluations) * 100, 1) if total_evaluations else 0,
        }
        for row in qs.values('establishment__governorate__name_ar').annotate(total=Count('id')).order_by('-total')[:6]
    ]

    reference_data = _get_reference_data()
    classification_choices = Evaluation.CLASSIFICATION_CHOICES

    paginator = Paginator(qs, PAGE_SIZE)
    page_obj = paginator.get_page(request.GET.get('page'))

    return render(request, 'inspections/evaluation_list.html', {
        'evaluations': page_obj,
        'page_obj': page_obj,
        'governorates': [
            {'governorate__id': governorate['id'], 'governorate__name_ar': governorate['name_ar']}
            for governorate in reference_data['governorates']
        ],
        'wilayats': [
            {
                'wilayat__id': wilayat['id'],
                'wilayat__name_ar': wilayat['name_ar'],
                'governorate__id': wilayat['governorate_id'],
            }
            for wilayat in reference_data['wilayats']
        ],
        'classification_choices': classification_choices,
        'activity_options': activity_options,
        'selected_governorate': governorate_id,
        'selected_wilayat': wilayat_id,
        'selected_classification': classification,
        'selected_activity': activity,
        'q': q,
        'total_evaluations': total_evaluations,
        'completed_evaluations_count': approval_counts.get('completed', 0),
        'draft_evaluations_count': approval_counts.get('draft', 0),
        'avg_evaluation_percentage': qs.aggregate(avg=Avg('percentage'))['avg'] or 0,
        'evaluated_establishments_count': qs.values('establishment_id').distinct().count(),
        'evaluation_classification_chart': evaluation_classification_chart,
        'evaluation_governorate_chart': evaluation_governorate_chart,
        'hide_establishments_nav': True,
        'hide_water_nav': True,
        'show_reports_nav': True,
    })


def create_evaluation_items(evaluation):
    criteria = list(Criterion.objects.filter(is_active=True).only('id', 'weight'))
    EvaluationItem.objects.bulk_create(
        [
            EvaluationItem(
                evaluation=evaluation,
                criterion=criterion,
                status='compliant',
                score_awarded=criterion.weight,
            )
            for criterion in criteria
        ],
        ignore_conflicts=True,
    )

    records = list(RequiredRecord.objects.filter(is_active=True).only('id'))
    EvaluationRecordCheck.objects.bulk_create(
        [
            EvaluationRecordCheck(evaluation=evaluation, record=record)
            for record in records
        ],
        ignore_conflicts=True,
    )


def sync_evaluation_items_with_active_template(evaluation):
    create_evaluation_items(evaluation)


@login_required
def evaluation_create(request):
    profile = _get_profile(request.user)
    establishments_qs = Establishment.objects.all()
    establishments_qs, _ = _apply_rbac(establishments_qs, None, profile)
    default_establishment = establishments_qs.order_by('commercial_name', 'id').first()

    if not default_establishment:
        messages.error(request, 'لا توجد منشآت متاحة لإنشاء تقييم. يرجى إضافة منشأة أولاً.')
        return redirect('/establishments/new/?next=/evaluations/new/')

    if request.method == 'POST':
        # Handle form submission
        establishment_id = request.POST.get('establishment')
        visit_date = request.POST.get('visit_date')
        notes = request.POST.get('notes', '')
        
        try:
            establishment = establishments_qs.get(id=establishment_id)
        except (Establishment.DoesNotExist, TypeError, ValueError):
            messages.error(request, 'المنشأة المختارة غير موجودة.')
            return redirect('evaluation_create')
        
        # Create evaluation
        evaluation = Evaluation.objects.create(
            establishment=establishment,
            inspector=request.user,
            visit_date=visit_date or timezone.localdate(),
            notes=notes,
            iso_22000_certificate=request.POST.get('iso_22000_certificate', '').strip(),
            haccp_certificate=request.POST.get('haccp_certificate', '').strip(),
            other_quality_certificate=request.POST.get('other_quality_certificate', '').strip(),
        )
        create_evaluation_items(evaluation)
        
        # Process evaluation items
        all_items = EvaluationItem.objects.filter(evaluation=evaluation).select_related('criterion')
        for item in all_items:
            status = request.POST.get(f'item_{item.criterion_id}', 'compliant')
            item.status = status
            item.remarks = request.POST.get(f'remarks_{item.criterion_id}', '')
            item.corrective_action = request.POST.get(f'corrective_{item.criterion_id}', '')
            item.save()
            
            # Handle image uploads
            image_file = request.FILES.get(f'image_{item.criterion_id}')
            if image_file:
                EvaluationImage.objects.create(
                    evaluation=evaluation,
                    criterion=item.criterion,
                    image=image_file,
                )
        
        record_checks = list(
            EvaluationRecordCheck.objects.filter(evaluation=evaluation).select_related('record')
        )
        for record_check in record_checks:
            record_check.is_available = request.POST.get(f'record_{record_check.record_id}') == 'on'
            record_check.remarks = request.POST.get(f'record_remarks_{record_check.record_id}', '').strip()
        EvaluationRecordCheck.objects.bulk_update(record_checks, ['is_available', 'remarks'])

        # Calculate results
        evaluation.calculate_results(items=list(all_items), record_checks=record_checks)
        
        # Check if save as draft
        if 'save_as_draft' in request.POST:
            evaluation.approval_status = 'draft'
            evaluation.save(update_fields=['total_points', 'percentage', 'classification', 'approval_status'])
            messages.success(request, 'تم حفظ التقييم كمسودة.')
        else:
            evaluation.approval_status = 'completed'
            evaluation.save(update_fields=['total_points', 'percentage', 'classification', 'approval_status'])
            messages.success(request, f'تم اعتماد التقييم بنجاح. النسبة: {evaluation.percentage}% - التصنيف: {evaluation.get_classification_display()}')
        
        EvaluationActivityLog.objects.create(
            evaluation=evaluation,
            user=request.user,
            action='إنشاء وحفظ تقييم',
            notes='تم إنشاء التقييم من استمارة التقييم المتقدمة.',
        )
        return redirect('evaluation_list')
    
    # GET request - show form
    from django.forms import ModelForm, Form
    from django import forms as dj_forms
    
    class SimpleEvaluationForm(ModelForm):
        establishment = dj_forms.ModelChoiceField(
            queryset=establishments_qs,
            label='المنشأة',
            widget=dj_forms.Select(attrs={'class': 'form-control'})
        )
        
        class Meta:
            model = Evaluation
            fields = ['establishment', 'visit_date', 'notes']
            labels = {
                'visit_date': 'تاريخ التقييم',
                'notes': 'الملاحظات العامة',
            }
    
    form = SimpleEvaluationForm()
    
    # Get sections and criteria
    sections = EvaluationSection.objects.filter(
        criteria__is_active=True
    ).distinct().prefetch_related(
        Prefetch(
            'criteria',
            Criterion.objects.filter(is_active=True).order_by('sort_order', 'code')
        )
    ).order_by('sort_order', 'name_ar')
    records = RequiredRecord.objects.filter(is_active=True).order_by('name_ar', 'id')
    
    return render(request, 'inspections/evaluation_create.html', {
        'form': form,
        'establishment': default_establishment,
        'sections': sections,
        'records': records,
    })


@login_required
def evaluation_update(request, pk):
    profile = _get_profile(request.user)
    evaluation = get_object_or_404(
        _get_allowed_evaluations(
            request.user,
            Evaluation.objects.select_related('establishment', 'inspector', 'reviewer').prefetch_related('team_members'),
        ),
        pk=pk,
    )

    # المفتش لا يستطيع تعديل تقييمات الآخرين
    if profile and profile.role == 'inspector' and evaluation.inspector != request.user:
        messages.error(request, 'ليس لديك صلاحية لتعديل هذا التقييم.')
        return redirect('evaluation_list')

    sync_evaluation_items_with_active_template(evaluation)

    queryset = EvaluationItem.objects.filter(
        evaluation=evaluation,
        criterion__is_active=True,
    ).select_related(
        'criterion', 'criterion__section'
    ).order_by(
        'criterion__section__sort_order', 'criterion__sort_order', 'id'
    )

    EvaluationItemFormSet = modelformset_factory(
        EvaluationItem,
        form=EvaluationItemForm,
        extra=0,
        can_delete=False,
    )
    EvaluationRecordFormSet = modelformset_factory(
        EvaluationRecordCheck,
        form=EvaluationRecordCheckForm,
        extra=0,
        can_delete=False,
    )
    EvaluationTeamMemberFormSet = modelformset_factory(
        EvaluationTeamMember,
        form=EvaluationTeamMemberForm,
        extra=1,
        can_delete=True,
    )
    record_queryset = EvaluationRecordCheck.objects.filter(
        evaluation=evaluation,
        record__is_active=True,
    ).select_related('record').order_by('record__name_ar', 'id')
    team_queryset = EvaluationTeamMember.objects.filter(
        evaluation=evaluation
    ).order_by('sort_order', 'id')
    meta_form = EvaluationHeaderForm(request.POST or None, instance=evaluation, prefix='meta')

    if request.method == 'POST':
        formset = EvaluationItemFormSet(request.POST, queryset=queryset)
        record_formset = EvaluationRecordFormSet(request.POST, queryset=record_queryset, prefix='records')
        team_formset = EvaluationTeamMemberFormSet(request.POST, queryset=team_queryset, prefix='team')
        if meta_form.is_valid() and formset.is_valid() and record_formset.is_valid() and team_formset.is_valid():
            from django.db import transaction as _tx
            with _tx.atomic():
                meta_form.save()
                formset.save()
                record_formset.save()
                team_members = team_formset.save(commit=False)
                for deleted_member in team_formset.deleted_objects:
                    deleted_member.delete()
                for index, member in enumerate(team_members, start=1):
                    member.evaluation = evaluation
                    member.sort_order = index
                    member.save()

                # استخدام البيانات الموجودة في الذاكرة بدل إعادة الجلب من DB
                all_items = [f.instance for f in formset.forms if f.instance.pk]
                all_records = [f.instance for f in record_formset.forms if f.instance.pk]

                for item in all_items:
                    if item.status == 'non_compliant' and not (item.corrective_action or '').strip():
                        item.corrective_action = _build_default_corrective_action(item)
                        item.save(update_fields=['corrective_action'])

                # معالجة الصور لجميع البنود (وليس فقط المعدّلة)
                for form in formset.forms:
                    form_item = form.instance
                    if not form_item.pk:
                        continue
                    image_field = request.FILES.get(f'image_{form_item.id}')
                    caption = request.POST.get(f'caption_{form_item.id}', '').strip()
                    if image_field:
                        EvaluationImage.objects.create(
                            evaluation=evaluation,
                            criterion=form_item.criterion,
                            image=image_field,
                            caption=caption,
                        )

                # تمرير البيانات الموجودة في الذاكرة لتجنب استعلامات DB إضافية
                _sync_corrective_actions_for_evaluation(evaluation, request.user, pre_loaded_items=all_items)
                evaluation.calculate_results(items=all_items, record_checks=all_records)
                evaluation.approval_status = 'completed'
                evaluation.save(update_fields=['total_points', 'percentage', 'classification', 'approval_status'])
                EvaluationActivityLog.objects.create(evaluation=evaluation, user=request.user, action='إنهاء التقييم', notes='تم حفظ بنود التقييم وإنهاء التقييم.')

                # محرك التأهيل الذكي وربط البنود غير المستوفية وخطة HACCP
                from inspections.models import QualificationFollowUp, HACCPFile
                pct = float(evaluation.percentage)
                if pct >= 70:
                    qual_status = 'completed'
                elif pct >= 41:
                    qual_status = 'in_progress'
                else:
                    qual_status = 'stalled'

                qf, created = QualificationFollowUp.objects.get_or_create(
                    establishment=evaluation.establishment,
                    defaults={
                        'governorate': evaluation.establishment.governorate.name_ar,
                        'establishment_name': evaluation.establishment.commercial_name,
                        'activity_type': evaluation.establishment.activity_type,
                        'current_status': qual_status,
                        'evaluation': evaluation,
                    }
                )
                if not created:
                    qf.current_status = qual_status
                    qf.evaluation = evaluation
                    qf.save(update_fields=['current_status', 'evaluation'])

                # مزامنة ملاحظات ملفات HACCP الموجودة فقط (حقل الملف إلزامي ولا يمكن إنشاء سجل بدون ملف)
                for item in all_items:
                    if item.status == 'non_compliant':
                        criterion_text = item.criterion.text_ar
                        if 'تتبع' in criterion_text or 'تحليل مخاطر' in criterion_text or 'سجلات' in criterion_text or 'نظافة' in criterion_text:
                            haccp_file_type = 'prps' if 'نظافة' in criterion_text else (
                                'traceability' if 'تتبع' in criterion_text else (
                                    'records' if 'سجلات' in criterion_text else 'hazard_analysis'
                                )
                            )
                            related_haccp_file = HACCPFile.objects.filter(
                                establishment=evaluation.establishment,
                                file_type=haccp_file_type,
                                title=criterion_text,
                            ).first()
                            if related_haccp_file:
                                auto_note = f'تم ربطه تلقائيًا بناءً على بند غير مستوفي في التقييم رقم {evaluation.pk}'
                                if auto_note not in (related_haccp_file.notes or ''):
                                    related_haccp_file.notes = '\n'.join(
                                        part for part in [related_haccp_file.notes.strip(), auto_note] if part
                                    )
                                    related_haccp_file.save(update_fields=['notes'])

            # مسح كاش PDF بعد الحفظ حتى لا يُعرض تقرير قديم
            _clear_evaluation_pdf_cache(evaluation.pk)
            messages.success(request, f'تم حفظ وإنهاء التقييم بنجاح. النسبة: {evaluation.percentage}% - التصنيف: {evaluation.get_classification_display()}')
            return redirect('evaluation_update', pk=evaluation.pk)
    else:
        formset = EvaluationItemFormSet(queryset=queryset)
        record_formset = EvaluationRecordFormSet(queryset=record_queryset, prefix='records')
        team_formset = EvaluationTeamMemberFormSet(queryset=team_queryset, prefix='team')

    grouped_forms = OrderedDict()
    status_counts = {
        row['status']: row['total']
        for row in queryset.values('status').annotate(total=Count('id'))
    }
    non_compliant_count = status_counts.get('non_compliant', 0)
    compliant_count = status_counts.get('compliant', 0)
    has_evaluation_result = EvaluationActivityLog.objects.filter(
        evaluation=evaluation,
        action__in=['حفظ التقييم', 'إنهاء التقييم'],
    ).exists()
    high_risk_non_compliant = list(
        queryset.filter(
            status='non_compliant',
            criterion__risk_level__in=['high', 'critical'],
        ).values_list('criterion__code', flat=True)
    )

    image_map = {}
    for image in EvaluationImage.objects.filter(evaluation=evaluation).select_related('criterion').order_by('id'):
        image_map.setdefault(image.criterion_id, []).append(image)

    for form in formset.forms:
        section = form.instance.criterion.section
        grouped_forms.setdefault(section, [])
        grouped_forms[section].append(form)

    signature_rows = [
        {
            'name': getattr(evaluation.inspector, 'get_full_name', lambda: '')() or evaluation.inspector.username,
            'job_title': 'المفتش / المقيم',
        }
    ]
    signature_rows.extend(
        {
            'name': member.full_name,
            'job_title': member.job_title,
        }
        for member in evaluation.team_members.all()
    )
    if evaluation.reviewer:
        signature_rows.append({
            'name': getattr(evaluation.reviewer, 'get_full_name', lambda: '')() or evaluation.reviewer.username,
            'job_title': 'المراجع الفني',
        })
    while len(signature_rows) < 4:
        signature_rows.append({'name': '', 'job_title': ''})

    context = {
        'evaluation': evaluation,
        'meta_form': meta_form,
        'formset': formset,
        'record_formset': record_formset,
        'team_formset': team_formset,
        'grouped_forms': grouped_forms,
        'non_compliant_count': non_compliant_count,
        'compliant_count': compliant_count,
        'has_evaluation_result': has_evaluation_result,
        'high_risk_non_compliant': high_risk_non_compliant,
        'image_map': image_map,
    }
    return render(request, 'inspections/evaluation_update.html', context)


@login_required
def evaluation_submit(request, pk):
    profile = _get_profile(request.user)
    evaluation = get_object_or_404(_get_allowed_evaluations(request.user), pk=pk)

    # المفتش لا يستطيع إنهاء تقييمات الآخرين
    if profile and profile.role == 'inspector' and evaluation.inspector != request.user:
        messages.error(request, 'ليس لديك صلاحية لإنهاء هذا التقييم.')
        return redirect('evaluation_list')

    evaluation.mark_completed()
    _clear_evaluation_pdf_cache(evaluation.pk)
    EvaluationActivityLog.objects.create(evaluation=evaluation, user=request.user, action='إنهاء التقييم')
    messages.success(request, 'تم إنهاء التقييم وحفظه كحالة مكتملة.')
    return redirect('evaluation_list')


@login_required
def corrective_action_list(request):
    qs = (
        CorrectiveActionLog.objects.select_related('evaluation', 'criterion')
        .order_by('-id')
    )
    paginator = Paginator(qs, PAGE_SIZE)
    page_obj = paginator.get_page(request.GET.get('page'))
    return render(
        request,
        'inspections/corrective_list.html',
        {
            'actions': page_obj,
            'page_obj': page_obj,
        },
    )


@login_required
def corrective_action_create(request):
    form = CorrectiveActionForm(request.POST or None)
    if form.is_valid():
        obj = form.save(commit=False)
        obj.created_by = request.user
        obj.save()
        messages.success(request, 'تم حفظ الإجراء التصحيحي.')
        return redirect('corrective_action_list')
    return render(request, 'inspections/form.html', {'form': form, 'title': 'إضافة إجراء تصحيحي'})


@login_required
def corrective_action_update(request, pk):
    corrective_action = get_object_or_404(CorrectiveActionLog, pk=pk)
    form = CorrectiveActionForm(request.POST or None, instance=corrective_action)
    if form.is_valid():
        form.save()
        messages.success(request, 'تم تحديث الإجراء التصحيحي.')
        return redirect('corrective_action_list')
    return render(request, 'inspections/form.html', {'form': form, 'title': 'تعديل إجراء تصحيحي'})


@login_required
def export_establishments_excel(request):
    profile = _get_profile(request.user)
    qs = Establishment.objects.select_related('governorate', 'wilayat').all()
    qs, _ = _apply_rbac(qs, None, profile)

    governorate_id = request.GET.get('governorate', '').strip()
    wilayat_id = request.GET.get('wilayat', '').strip()
    activity = request.GET.get('activity', '').strip()
    q = request.GET.get('q', '').strip()
    normalized_q = _normalize_digit_text(q)

    if governorate_id:
        qs = qs.filter(governorate_id=governorate_id)
    if wilayat_id:
        qs = qs.filter(wilayat_id=wilayat_id)
    if activity:
        qs = qs.filter(activity_type=activity)
    if q:
        filters = (
            Q(commercial_name__icontains=q) |
            Q(activity_type__icontains=q) |
            Q(license_no__icontains=q) |
            Q(commercial_reg__icontains=q)
        )
        if normalized_q.isdigit():
            filters |= Q(establishment_no=int(normalized_q))
        qs = qs.filter(filters)

    wb = Workbook()
    ws = wb.active
    ws.title = 'المنشآت'
    ws.append(['رقم المنشأة', 'الاسم التجاري', 'النشاط', 'المحافظة', 'الولاية', 'رقم الترخيص', 'السجل التجاري', 'الحالة'])
    for e in qs:
        ws.append([
            e.establishment_no,
            e.commercial_name, e.activity_type, e.governorate.name_ar, e.wilayat.name_ar,
            e.license_no, e.commercial_reg, e.get_status_display()
        ])
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="establishments.xlsx"'
    wb.save(response)
    return response


@login_required
def export_qualification_followups_excel(request):
    qs = _with_qualification_visit_no(
        QualificationFollowUp.objects.select_related('establishment').all()
    )

    q = request.GET.get('q', '').strip()
    governorate = request.GET.get('governorate', '').strip()
    status = request.GET.get('status', '').strip()
    activity = request.GET.get('activity', '').strip()
    if q:
        normalized_q = _normalize_digit_text(q)
        filters = (
            Q(establishment_name__icontains=q) |
            Q(activity_type__icontains=q) |
            Q(governorate__icontains=q) |
            Q(establishment__license_no__icontains=q)
        )
        if normalized_q.isdigit():
            filters |= Q(establishment__establishment_no=int(normalized_q))
        qs = qs.filter(filters)
    if governorate:
        qs = qs.filter(governorate=governorate)
    if status:
        qs = qs.filter(current_status=status)
    if activity:
        qs = qs.filter(activity_type=activity)

    wb = Workbook()
    ws = wb.active
    ws.title = 'متابعة التأهيل'
    ws.append([
        'م', 'المحافظة', 'اسم المنشأة', 'نوع النشاط', 'الحالة الحالية',
        'أنظمة الجودة وسلامة الغذاء', 'تاريخ البدء', 'تاريخ الإنجاز المتوقع',
        'نسبة الإنجاز (%)', 'التحديات', 'ملاحظات', 'رقم المنشأة المرجعي',
        'رقم الزيارة المرجعي', 'سنة الزيارة', 'رقم الزيارة لنفس المنشأة',
        'رمز المحافظة', 'رمز النشاط', 'مفتاح الربط Django',
    ])
    for index, item in enumerate(qs, start=1):
        ws.append([
            index,
            item.governorate,
            item.establishment_name,
            item.activity_type,
            item.get_current_status_display(),
            item.custom_quality_system or item.quality_system,
            item.start_date,
            item.expected_completion_date,
            item.progress_percent,
            item.challenges,
            item.notes,
            item.facility_reference_code,
            item.visit_reference_code,
            item.visit_year,
            item.visit_no,
            item.governorate_code,
            item.activity_code,
            item.django_link_key,
        ])
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="qualification_followups.xlsx"'
    wb.save(response)
    return response


@login_required
def export_evaluations_excel(request):
    profile = _get_profile(request.user)
    qs = Evaluation.objects.select_related(
        'establishment', 'establishment__governorate', 'establishment__wilayat', 'inspector'
    ).prefetch_related(
        Prefetch('items', queryset=EvaluationItem.objects.filter(status='non_compliant', criterion__is_active=True).select_related('criterion')),
        Prefetch('record_checks', queryset=EvaluationRecordCheck.objects.filter(is_available=False, record__is_active=True).select_related('record')),
    ).all()
    _, qs = _apply_rbac(None, qs, profile)
    governorate_id = request.GET.get('governorate', '').strip()
    wilayat_id = request.GET.get('wilayat', '').strip()
    classification = request.GET.get('classification', '').strip()
    activity = request.GET.get('activity', '').strip()
    q = request.GET.get('q', '').strip()
    normalized_q = _normalize_digit_text(q)

    if governorate_id:
        qs = qs.filter(establishment__governorate_id=governorate_id)
    if wilayat_id:
        qs = qs.filter(establishment__wilayat_id=wilayat_id)
    if classification:
        qs = qs.filter(classification=classification)
    if activity:
        qs = qs.filter(establishment__activity_type=activity)
    if q:
        filters = (
            Q(establishment__commercial_name__icontains=q) |
            Q(establishment__activity_type__icontains=q) |
            Q(establishment__license_no__icontains=q) |
            Q(establishment__commercial_reg__icontains=q)
        )
        if normalized_q.isdigit():
            filters |= Q(establishment__establishment_no=int(normalized_q))
        qs = qs.filter(filters)

    wb = Workbook()
    ws = wb.active
    ws.title = 'تقارير التقييم'
    ws.append([
        'مرجع التقرير', 'رقم المنشأة', 'اسم المنشأة', 'المحافظة', 'الولاية', 'رقم الترخيص', 'تاريخ الزيارة',
        'المفتش', 'النسبة', 'التصنيف', 'الحالة', 'البنود غير المستوفية',
    ])

    for evaluation in qs:
        non_compliant = evaluation.items.all()
        missing_records = evaluation.record_checks.all()
        non_compliant_text = ' | '.join(
            f'{item.criterion.code}. {item.criterion.text_ar}' + (f' - {item.remarks}' if item.remarks else '')
            for item in non_compliant
        )
        missing_records_text = ' | '.join(record_check.record.name_ar for record_check in missing_records)
        issues_text = ' | '.join(filter(None, [non_compliant_text, missing_records_text])) or 'لا توجد بنود غير مستوفية'
        ws.append([
            evaluation.report_reference_no,
            evaluation.establishment.establishment_no,
            evaluation.establishment.commercial_name,
            evaluation.establishment.governorate.name_ar,
            evaluation.establishment.wilayat.name_ar,
            evaluation.establishment.license_no,
            str(evaluation.visit_date),
            evaluation.inspector.get_full_name() or evaluation.inspector.username,
            float(evaluation.percentage),
            evaluation.get_classification_display(),
            evaluation.get_approval_status_display(),
            issues_text,
        ])

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="evaluation_reports.xlsx"'
    wb.save(response)
    return response


_static_path_cache: dict = {}


def link_callback(uri, rel):
    if uri.startswith(settings.MEDIA_URL):
        path = os.path.join(settings.MEDIA_ROOT, uri.replace(settings.MEDIA_URL, ''))
    elif uri.startswith(settings.STATIC_URL):
        relative = uri.replace(settings.STATIC_URL, '')
        if relative not in _static_path_cache:
            _static_path_cache[relative] = finders.find(relative)
        path = _static_path_cache[relative]
    else:
        return uri

    if not path or not os.path.isfile(path):
        raise Exception(f'تعذر العثور على الملف: {uri}')
    return path



def render_to_pdf(template_src, context_dict):
    template = get_template(template_src)
    html = template.render(context_dict)
    result = BytesIO()
    original_named_temp_file = pisa_files.tempfile.NamedTemporaryFile

    def reopenable_named_temp_file(*args, **kwargs):
        kwargs.setdefault('delete', False)
        return original_named_temp_file(*args, **kwargs)

    pisa_files.tempfile.NamedTemporaryFile = reopenable_named_temp_file
    try:
        pdf = pisa.pisaDocument(
            BytesIO(html.encode('UTF-8')),
            result,
            encoding='UTF-8',
            link_callback=link_callback,
        )
    finally:
        pisa_files.tempfile.NamedTemporaryFile = original_named_temp_file

    if not pdf.err:
        return HttpResponse(result.getvalue(), content_type='application/pdf')
    return None


def get_pdf_font_context():
    return {}


def _build_evaluation_report_context(evaluation):
    items_list = list(
        EvaluationItem.objects.filter(evaluation=evaluation, status='non_compliant', criterion__is_active=True)
        .select_related('criterion', 'criterion__section')
        .order_by('criterion__section__sort_order', 'criterion__sort_order', 'id')
    )

    grouped_items = OrderedDict()
    record_checks = list(
        EvaluationRecordCheck.objects.filter(evaluation=evaluation, record__is_active=True)
        .select_related('record')
        .order_by('record__name_ar', 'id')
    )
    images_by_criterion = {}
    for image in EvaluationImage.objects.filter(evaluation=evaluation).select_related('criterion').order_by('id'):
        images_by_criterion.setdefault(image.criterion_id, []).append(image)

    for item in items_list:
        section = item.criterion.section
        grouped_items.setdefault(section, [])
        grouped_items[section].append(item)

    # النتائج محفوظة مسبقاً - لا حاجة لإعادة الحساب عند كل فتح للتقرير

    def _person_name(user):
        if not user:
            return ''
        profile_name = getattr(getattr(user, 'userprofile', None), 'full_name', '')
        if (profile_name or '').strip():
            return profile_name.strip()
        full_name = getattr(user, 'get_full_name', lambda: '')() or ''
        return full_name.strip()

    signature_rows = []
    signature_rows.extend(
        {
            'name': member.full_name,
            'job_title': member.job_title,
        }
        for member in evaluation.team_members.all()
        if (member.full_name or '').strip() or (member.job_title or '').strip()
    )
    if evaluation.reviewer:
        signature_rows.append({
            'name': _person_name(evaluation.reviewer),
            'job_title': 'المراجع الفني',
        })

    return {
        'evaluation': evaluation,
        'grouped_items': grouped_items,
        'images_by_criterion': images_by_criterion,
        'record_checks': record_checks,
        'signature_rows': signature_rows,
        'non_compliant_total': len(items_list),
    }


def _set_docx_rtl(paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = alignment
    p_pr = paragraph._element.get_or_add_pPr()
    if p_pr.find(qn('w:bidi')) is None:
        p_pr.append(OxmlElement('w:bidi'))


def _set_table_rtl(table):
    """اضبط الجدول ليكون من اليمين إلى اليسار."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    if tblPr.find(qn('w:bidiVisual')) is None:
        bidi = OxmlElement('w:bidiVisual')
        tblPr.append(bidi)


def _set_docx_table_column_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            if index >= len(row.cells):
                continue
            cell = row.cells[index]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(int(width.inches * 1440)))
            tc_w.set(qn('w:type'), 'dxa')


def _shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def _set_cell_text(
    cell,
    text,
    bold=False,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
    fill=None,
    font_color=None,
):
    cell.vertical_alignment = vertical_alignment
    if fill:
        _shade_cell(cell, fill)
    cell.text = ''
    paragraph = cell.paragraphs[0]
    _set_docx_rtl(paragraph, alignment=alignment)
    run = paragraph.add_run(str(text or ''))
    run.bold = bold
    run.font.name = 'Tahoma'
    run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
    run.font.size = Pt(10)
    if font_color:
        run.font.color.rgb = RGBColor.from_string(font_color)


def _add_docx_heading(document, text, level=1):
    paragraph = document.add_heading('', level=level)
    _set_docx_rtl(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run(str(text or ''))
    run.font.name = 'Tahoma'
    run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
    return paragraph


def _add_docx_paragraph(document, label, value=''):
    paragraph = document.add_paragraph()
    _set_docx_rtl(paragraph)
    label_run = paragraph.add_run(str(label or ''))
    label_run.bold = True
    label_run.font.name = 'Tahoma'
    label_run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
    value_run = paragraph.add_run(str(value or ''))
    value_run.font.name = 'Tahoma'
    value_run._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
    return paragraph


def _format_docx_date(value):
    if not value:
        return ''
    return value.strftime('%Y/%m/%d')


def _format_docx_datetime(value):
    if not value:
        return ''
    return value.strftime('%Y/%m/%d %H:%M')


def _format_docx_file(value):
    if not value:
        return 'غير مرفق'
    name = getattr(value, 'name', '') or ''
    return name.split('/')[-1] if name else 'مرفق'


def _add_docx_label_value_table(document, rows, column_widths=None):
    table = document.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    _set_table_rtl(table)
    if column_widths:
        _set_docx_table_column_widths(table, column_widths)

    for row_values in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], row_values[0], bold=True, fill='E8F3ED')
        _set_cell_text(cells[1], row_values[1] or '-', alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(cells[2], row_values[2], bold=True, fill='E8F3ED')
        _set_cell_text(cells[3], row_values[3] or '-', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if column_widths:
        _set_docx_table_column_widths(table, column_widths)
    return table


def _add_docx_report_header(section):
    header_path = finders.find('images/report_header.png')
    if not header_path or not os.path.exists(header_path):
        return
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    _set_docx_rtl(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    available_width = section.page_width - section.left_margin - section.right_margin
    try:
        paragraph.add_run().add_picture(header_path, width=available_width)
    except Exception:
        return


def _build_evaluation_docx(evaluation):
    context = _build_evaluation_report_context(evaluation)
    grouped_items = context['grouped_items']
    images_by_criterion = context['images_by_criterion']
    record_checks = context['record_checks']

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(1.65)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.15)
    _add_docx_report_header(section)

    styles = document.styles
    styles['Normal'].font.name = 'Tahoma'
    styles['Normal']._element.rPr.rFonts.set(qn('w:cs'), 'Tahoma')
    styles['Normal'].font.size = Pt(10)

    establishment = evaluation.establishment

    _add_docx_heading(document, f'تقرير زيارة ميدانية إلى شركة: {establishment.commercial_name}', 1)

    summary_rows = [
        ('مرجع التقرير', evaluation.report_reference_no, 'تاريخ الزيارة', _format_docx_date(evaluation.visit_date)),
        ('اسم المنشأة', establishment.commercial_name, 'رقم المنشأة', establishment.establishment_no),
        ('المحافظة', establishment.governorate.name_ar, 'الولاية', establishment.wilayat.name_ar),
        ('نسبة الامتثال', f'{evaluation.percentage}%', 'التصنيف', evaluation.get_classification_display()),
        ('المفتش / المقيم', evaluation.inspector.get_full_name() or evaluation.inspector.username, 'حالة التقييم', evaluation.get_approval_status_display()),
    ]
    _add_docx_label_value_table(document, summary_rows, [Inches(1.45), Inches(3.5), Inches(1.45), Inches(3.5)])

    document.add_paragraph()
    _add_docx_heading(document, 'بيانات المنشأة التفصيلية', 2)
    establishment_rows = [
        ('الرقم المرجعي', establishment.reference_no, 'رقم المنشأة', establishment.establishment_no),
        ('الاسم التجاري', establishment.commercial_name, 'حالة المنشأة', establishment.get_status_display()),
        ('النشاط الرئيسي', establishment.activity_type, 'رقم رخصة النشاط', establishment.license_no),
        ('رقم السجل التجاري', establishment.commercial_reg, 'المحافظة', establishment.governorate.name_ar),
        ('الولاية', establishment.wilayat.name_ar, 'مدير الجودة / سلامة الغذاء', establishment.manager_name),
        ('رقم التواصل', establishment.contact_phone, 'البريد الإلكتروني', establishment.contact_email),
        ('عدد الموظفين', establishment.employee_count, 'الطاقة الإنتاجية', establishment.production_capacity),
        ('نوع المنتجات', establishment.product_types, 'الموقع المباشر', establishment.direct_location_url),
        ('خط العرض', establishment.latitude, 'خط الطول', establishment.longitude),
        ('تاريخ الإنشاء', _format_docx_datetime(establishment.created_at), 'تاريخ التحديث', _format_docx_datetime(establishment.updated_at)),
        ('السجل التجاري', _format_docx_file(establishment.doc_commercial_register), 'الترخيص البلدي', _format_docx_file(establishment.doc_municipal_license)),
        ('شهادات الجودة', _format_docx_file(establishment.doc_quality_certificates), 'مخططات المصنع', _format_docx_file(establishment.doc_factory_layout)),
    ]
    _add_docx_label_value_table(document, establishment_rows, [Inches(1.45), Inches(3.5), Inches(1.45), Inches(3.5)])

    document.add_paragraph()
    _add_docx_heading(document, 'دليل الهاسب للشركة إن وجد', 2)
    certificate_rows = [
        ('آيزو 22000', evaluation.iso_22000_certificate, 'الهاسب', evaluation.haccp_certificate),
        ('أخرى', evaluation.other_quality_certificate, '', ''),
    ]
    _add_docx_label_value_table(document, certificate_rows, [Inches(1.45), Inches(3.5), Inches(1.45), Inches(3.5)])
    records_table = document.add_table(rows=1, cols=3)
    records_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    records_table.style = 'Table Grid'
    _set_table_rtl(records_table)
    for index, header in enumerate(['المتطلبات', 'وجود المستند', 'الملاحظات']):
        _set_cell_text(records_table.rows[0].cells[index], header, bold=True, fill='2F855A', font_color='FFFFFF')
    for record_check in record_checks:
        cells = records_table.add_row().cells
        _set_cell_text(cells[0], record_check.record.name_ar)
        _set_cell_text(cells[1], 'نعم' if record_check.is_available else 'لا')
        _set_cell_text(cells[2], record_check.remarks or '-')
    _set_docx_table_column_widths(records_table, [Inches(5), Inches(1.5), Inches(3)])

    document.add_paragraph()
    _add_docx_heading(document, 'البنود غير المستوفية والإجراءات التصحيحية', 2)

    if not grouped_items:
        _add_docx_paragraph(document, '', 'لا توجد بنود غير مستوفية.')
    else:
        for section_obj, items in grouped_items.items():
            _add_docx_heading(document, f'{section_obj.sort_order} - {section_obj.name_ar}', 3)
            has_remarks = any((item.remarks or '').strip() for item in items)
            num_cols = 5 if has_remarks else 4
            table = document.add_table(rows=1, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            _set_table_rtl(table)
            if has_remarks:
                headers = ['البند', 'نص البند غير المستوفي', 'الملاحظات', 'الإجراء التصحيحي', 'الصور']
                column_widths = [Inches(0.55), Inches(3.2), Inches(1.75), Inches(3.1), Inches(1.45)]
            else:
                headers = ['البند', 'نص البند غير المستوفي', 'الإجراء التصحيحي', 'الصور']
                column_widths = [Inches(0.55), Inches(4.15), Inches(3.7), Inches(1.65)]
            for index, header in enumerate(headers):
                _set_cell_text(table.rows[0].cells[index], header, bold=True, fill='2F855A', font_color='FFFFFF')
            _set_docx_table_column_widths(table, column_widths)
            for item in items:
                cells = table.add_row().cells
                _set_cell_text(cells[0], item.criterion.code)
                _set_cell_text(cells[1], item.criterion.text_ar)
                if has_remarks:
                    _set_cell_text(cells[2], item.remarks or '-')
                    _set_cell_text(cells[3], item.corrective_action or '-')
                    image_cell = cells[4]
                else:
                    _set_cell_text(cells[2], item.corrective_action or '-')
                    image_cell = cells[3]
                image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                image_cell.text = ''
                images = images_by_criterion.get(item.criterion_id, [])
                if images:
                    for image in images:
                        paragraph = image_cell.add_paragraph()
                        _set_docx_rtl(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                        image_path = getattr(image.image, 'path', '')
                        if image_path and os.path.exists(image_path):
                            try:
                                paragraph.add_run().add_picture(image_path, width=Inches(1.55))
                            except Exception:
                                paragraph.add_run('تعذر إدراج الصورة')
                        else:
                            paragraph.add_run('الصورة غير متوفرة')
                        if image.caption:
                            caption = image_cell.add_paragraph()
                            _set_docx_rtl(caption, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                            caption.add_run(image.caption)
                else:
                    _set_cell_text(image_cell, '-')
            _set_docx_table_column_widths(table, column_widths)

    # نتيجة التقييم (جدول)
    _add_docx_heading(document, 'نتيجة التقييم', 2)
    result_table = document.add_table(rows=1, cols=3)
    result_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    result_table.style = 'Table Grid'
    _set_table_rtl(result_table)
    headers = ['الوصف', 'النسبة', 'التصنيف']
    for index, header in enumerate(headers):
        _set_cell_text(result_table.rows[0].cells[index], header, bold=True, fill='2F855A', font_color='FFFFFF')
    _set_docx_table_column_widths(result_table, [Inches(5.5), Inches(2), Inches(2)])
    status = getattr(evaluation, 'establishment_status', None)
    if status:
        row = result_table.add_row().cells
        _set_cell_text(row[0], status.get('description', '') if isinstance(status, dict) else getattr(status, 'description', ''))
        _set_cell_text(row[1], f"{evaluation.percentage}%")
        _set_cell_text(row[2], status.get('label', '') if isinstance(status, dict) else getattr(status, 'label', ''))
        _set_docx_table_column_widths(result_table, [Inches(5.5), Inches(2), Inches(2)])

    # ملاحظات عامة (قبل فريق التقييم)
    if (evaluation.notes or '').strip():
        _add_docx_heading(document, 'ملاحظات عامة', 2)
        _add_docx_paragraph(document, '', evaluation.notes)

    # الإجراءات التصحيحية العامة (إن وجدت)
    if (evaluation.corrective_action or '').strip():
        _add_docx_heading(document, 'الإجراءات التصحيحية العامة', 2)
        _add_docx_paragraph(document, '', evaluation.corrective_action)

    # فريق التقييم والاعتماد
    if context['signature_rows']:
        _add_docx_heading(document, 'فريق التقييم والاعتماد', 2)
        sign_table = document.add_table(rows=1, cols=3)
        sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sign_table.style = 'Table Grid'
        _set_table_rtl(sign_table)
        for index, header in enumerate(['الاسم', 'المسمى الوظيفي', 'التوقيع']):
            _set_cell_text(sign_table.rows[0].cells[index], header, bold=True, fill='2F855A', font_color='FFFFFF')
        _set_docx_table_column_widths(sign_table, [Inches(3.2), Inches(3.2), Inches(3.1)])
        for signature in context['signature_rows']:
            cells = sign_table.add_row().cells
            _set_cell_text(cells[0], signature['name'])
            _set_cell_text(cells[1], signature['job_title'])
            _set_cell_text(cells[2], '')
        _set_docx_table_column_widths(sign_table, [Inches(3.2), Inches(3.2), Inches(3.1)])

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


@login_required
def evaluation_word(request, pk):
    evaluation = get_object_or_404(
        _get_allowed_evaluations(
            request.user,
            Evaluation.objects.select_related(
                'establishment',
                'establishment__governorate',
                'establishment__wilayat',
                'inspector',
                'reviewer',
            ).prefetch_related('team_members'),
        ),
        pk=pk,
    )
    docx_file = _build_evaluation_docx(evaluation)
    response = HttpResponse(
        docx_file.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="official_evaluation_{evaluation.id}.docx"'
    return response


PDF_CACHE_TIMEOUT = 60 * 10  # 10 دقائق


@login_required
def evaluation_pdf(request, pk):
    evaluation = get_object_or_404(
        _get_allowed_evaluations(
            request.user,
            Evaluation.objects.select_related('establishment', 'inspector', 'reviewer').prefetch_related('team_members'),
        ),
        pk=pk,
    )
    # استخدم الكاش فقط للتقييمات المكتملة (المسودات قد تتغير)
    pdf_cache_key = None
    if evaluation.approval_status == 'completed':
        pdf_cache_key = f'pdf_bytes:v3:eval:{pk}'
        cached_pdf = cache.get(pdf_cache_key)
        if cached_pdf is not None:
            response = HttpResponse(cached_pdf, content_type='application/pdf')
            response['Content-Disposition'] = f'inline; filename="official_evaluation_{pk}.pdf"'
            return response

    context = _build_evaluation_report_context(evaluation)
    response = render_to_pdf('inspections/evaluation_pdf_official.html', context)
    if response:
        if pdf_cache_key:
            cache.set(pdf_cache_key, response.content, PDF_CACHE_TIMEOUT)
        response['Content-Disposition'] = f'inline; filename="official_evaluation_{evaluation.id}.pdf"'
        return response
    return HttpResponse('حدث خطأ أثناء إنشاء ملف PDF')

